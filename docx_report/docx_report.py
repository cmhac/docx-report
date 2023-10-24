# -*- coding: utf-8 -*-
"""Generates a docx report"""

from datetime import datetime
import os
from typing import Optional
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH  # pylint: disable=no-name-in-module
import janitor  # pylint: disable=unused-import
import pandas as pd


class DocxReport:
    """Generates a docx report.

    Args:
        title (str): The title of the report.

    Attributes:
        _doc (docx.Document): The underlying docx document.
        title (str): The title of the report.

    """

    def __init__(self, title: str) -> None:
        """Initialize the DocxReport object.

        Args:
            title (str): The title of the report.

        """
        self._doc = docx.Document()
        self.title = title
        self._add_heading()

    def _add_heading(self) -> None:
        """Draws the heading for the report."""
        # Get the current time
        current_time = datetime.now().astimezone()

        # Format the time
        formatted_time = current_time.strftime("%B %-d, %-I:%M %p %Z")

        # Add subtitle above the header
        subtitle_p = self._doc.add_paragraph()
        subtitle_p.add_run(f"Generated on {formatted_time}")
        subtitle_p.style = "Subtitle"
        # Add the title
        self._doc.add_heading(self.title, 0)

    def add_paragraph(self, text: str) -> None:
        """Adds a paragraph to the document.

        Args:
            text (str): The text to add as a paragraph.

        """
        return self._doc.add_paragraph(text)

    def add_heading(self, text: str, level: int = 1) -> None:
        """Adds a heading to the document.

        Args:
            text (str): The text for the heading.
            level (int): The heading level (default is 1).

        """
        return self._doc.add_heading(text, level)

    def add_picture(self, filename: str, width: float = 5) -> None:
        """Adds a picture to the document.

        Args:
            filename (str): The path to the image file.
            width (float): The width of the picture in inches (default is 5).

        """
        return self._doc.add_picture(filename, width=docx.shared.Inches(width))

    @staticmethod
    def _cleanup_dataframe(
        df: pd.DataFrame,
        round_numeric: bool = True,
        round_decimals: int = 1,
        auto_format_dates: bool = True,
        rename_cols: Optional[dict] = None,
        strftime_format: str = "%Y-%m-%d",
    ) -> pd.DataFrame:
        """Cleans up a dataframe to be ready for plotting.

        Args:
            df (pd.DataFrame): The dataframe to clean up.
            round_numeric (bool): Whether to round numeric columns to
                2 decimal places (default is True).
            round_decimals (int): How many decimal places to round to (default is 1).
            auto_format_dates (bool): Whether to automatically format dates (default is True).
            rename_cols (Optional[dict]): A dictionary of columns to rename (default is None).
            strftime_format (str): The format to use when converting
                dates to strings (default is "%Y-%m-%d").

        Returns:
            pd.DataFrame: The cleaned up dataframe.
        """
        # rename the columns
        if rename_cols:
            df = df.rename(columns=rename_cols)

        # clean up the names
        df = df.clean_names().rename(columns=lambda x: x.replace("_", " "))

        # round all floats to 2 decimal places
        if round_numeric:
            df = df.applymap(
                lambda x: round(x, round_decimals) if isinstance(x, float) else x
            )

        # automatically clean up dates
        if auto_format_dates:
            # find all columns that are dates
            date_cols = df.select_dtypes(include="datetime").columns
            for date_col in date_cols:
                # convert to strings
                df[date_col] = df[date_col].dt.strftime(strftime_format)

        return df

    def _center_last_paragraph(self) -> None:
        """Centers the last paragraph in the doc."""
        last_paragraph = self._doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_plot(
        self,
        df: pd.DataFrame,
        title: str,
        x_label: str,
        y_label: str,
        rename_cols: Optional[dict] = None,
        **kwargs,
    ) -> None:
        """Uses matplotlib to plot a dataframe, then adds it to the docx file.

        Args:
            df (pd.DataFrame): The dataframe to plot.
            title (str): The title of the plot.
            x_label (str): The label for the x-axis.
            y_label (str): The label for the y-axis.
            rename_cols (Optional[dict]): A dictionary of columns to rename (default is None).
            **kwargs: Keyword arguments to pass to the pandas.DataFrame.plot function.
        """
        df = self._cleanup_dataframe(
            df, round_numeric=False, auto_format_dates=False, rename_cols=rename_cols
        )
        # create the plot
        ax = df.plot(**kwargs)
        ax.set_title(title)
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_label)
        # save the plot as a png
        ax.get_figure().savefig("temp.png")
        # add the plot to the docx file
        self._doc.add_picture("temp.png", width=docx.shared.Inches(5))
        # center the image
        self._center_last_paragraph()
        # delete the temp png
        os.remove("temp.png")

    def add_table(
        self,
        df: pd.DataFrame,
        include_index: bool = True,
        rename_cols: dict = None,
        pct_cols: list = None,
    ) -> None:
        """Turns a dataframe into a table in the document.

        Args:
            df (pd.DataFrame): The dataframe to turn into a table.
            include_index (bool): Whether to include the index as a column (default is True).
            rename_cols (dict): A dictionary of columns to rename (default is None).
            pct_cols (list): A list of columns to format as percentages (default is None).
        """
        # if include_index, reset the index so it's a column
        if include_index:
            df = df.reset_index()

        # save the original column names for later to check against pct_cols
        original_cols = df.columns.tolist()

        # do the cleanup
        df = self._cleanup_dataframe(df, rename_cols=rename_cols)

        # create the table based on the size of the dataframe
        rows = df.shape[0] + 1  # add 1 for the header
        cols = df.shape[1]
        table = self._doc.add_table(rows=rows, cols=cols)
        # set the style
        table.style = "TableGrid"

        # add the header
        header_cells = table.rows[0].cells
        for col_index, col in enumerate(df.columns):
            header_cells[col_index].text = col

        # add the data
        for row_index, row in df.iterrows():
            row_cells = table.rows[row_index + 1].cells
            for value_index, value in enumerate(row):
                # if value is numeric, add commas
                if isinstance(value, (int, float)):
                    # get the key from the original column names
                    # to check against pct_cols
                    key = original_cols[value_index]
                    # format percentages if needed
                    if pct_cols and key in pct_cols:
                        value = f"{value:.1%}"
                    else:
                        value = f"{value:,}"
                else:
                    value = str(value)
                row_cells[value_index].text = value
        return df

    def add_list_bullet(self, text: str) -> None:
        """Adds a bullet point to the document.

        Args:
            text (str): The text for the bullet point.

        """
        return self._doc.add_paragraph(text, style="List Bullet")

    def save(self, filename: str) -> None:
        """Saves the docx file.

        Args:
            filename (str): The filename to save the document as.

        """
        self._doc.save(filename)
