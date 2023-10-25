# -*- coding: utf-8 -*-
"""Tests the docx_report module."""

from datetime import datetime
import os
from docx.document import Document as DocumentBaseClass
from docx import Document
from docx.shared import Inches
import pandas as pd
import pytest
from docx_report import DocxReport

# pylint: disable=protected-access,redefined-outer-name


@pytest.fixture
def report():
    """Returns a DocxReport object"""
    return DocxReport("Test Report")


@pytest.fixture
def cleanup():
    """Removes the test.docx file after a test"""
    # 'yield' will return control to the test function and
    # resume here after the test function completes
    yield
    if os.path.exists("test.docx"):
        os.remove("test.docx")


# Sample data for testing
data = {
    "dates": [datetime(2021, 1, 1), datetime(2021, 1, 2), datetime(2021, 1, 3)],
    "values": [1.23, 4.56, 7.89],
}
df = pd.DataFrame(data)


def test_initialization(report):
    """Tests the initialization of a DocxReport object."""
    assert isinstance(report._doc, DocumentBaseClass)
    assert report.title == "Test Report"
    assert report._doc.paragraphs[1].text == "Test Report"


def test_cleanup_dataframe(report):
    """Tests the _cleanup_dataframe method."""
    cleaned_df = report._cleanup_dataframe(df)  # pylint: disable=protected-access
    assert cleaned_df["dates"].dtype == "object"  # Dates converted to strings
    assert cleaned_df["values"][0] == 1.2  # Values rounded to 1 decimal place


def test_cleanup_dataframe_rename_cols(report):
    """Tests the _cleanup_dataframe method with rename_cols."""
    initial_data = {"old_col1": [1, 2, 3], "old_col2": [4, 5, 6]}
    rename_dict = {"old_col1": "new_col1", "old_col2": "new_col2"}
    new_df = pd.DataFrame(initial_data)

    cleaned_df = report._cleanup_dataframe(new_df, rename_cols=rename_dict)

    assert list(cleaned_df.columns) == [
        "new col1",
        "new col2",
    ], f"Expected renamed columns, but got {list(cleaned_df.columns)}"


def test_add_paragraph(report):
    """Tests the add_paragraph method."""
    report.add_paragraph("Test Paragraph")

    paragraph = report._doc.paragraphs[-1]
    assert paragraph.text == "Test Paragraph"


def test_add_heading(report):
    """Tests the add_heading method."""
    report.add_heading("Test Heading", level=2)

    paragraph = report._doc.paragraphs[-1]
    assert paragraph.text == "Test Heading"
    assert paragraph.style.name == "Heading 2"


def test_add_picture(report):
    """Tests the add_picture method."""
    report.add_picture("test/test_data/test_image.png", width=6)

    last_shape = report._doc.inline_shapes[-1]
    assert last_shape.width == Inches(6)


def test_add_plot(tmp_path, report):
    """Tests the add_plot method."""
    report.add_plot(df, "Test Plot", "X Axis", "Y Axis")
    report.save(tmp_path / "test.docx")
    doc = Document(tmp_path / "test.docx")
    # check that an image was added
    assert len(doc.inline_shapes) == 1


def test_add_table(report):
    """Tests the add_table method."""
    report.add_table(df)
    assert report._doc.tables[0].rows[0].cells[0].text == "index"  # header added
    assert report._doc.tables[0].rows[1].cells[1].text == "2021-01-01"  # Data added


def doc_table_to_df(table):
    """Converts a docx table to a pandas DataFrame."""
    rows_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        rows_data.append(row_data)
    # Assuming the first row is headers
    return pd.DataFrame(rows_data[1:], columns=rows_data[0])


def test_add_table_pct_cols(
    cleanup,  # pylint: disable=unused-argument
    report,
):
    """Tests the add_table method with pct_cols."""
    initial_data = {"value": [0.1, 0.2, 0.3], "percent": [0.4, 0.5, 0.6]}
    new_df = pd.DataFrame(initial_data)

    report.add_table(new_df, pct_cols=["percent"])
    report.save("test.docx")

    doc = Document("test.docx")
    table = doc.tables[0]
    table_df = doc_table_to_df(table)
    expected_percent_col = ["40.0%", "50.0%", "60.0%"]
    assert (
        list(table_df["percent"]) == expected_percent_col
    ), f"Expected {expected_percent_col}, but got {list(table_df['percent'])}"


def test_add_list_bullet(report):
    """Tests the add_list_bullet method."""
    report.add_list_bullet("Test Bullet")
    assert report._doc.paragraphs[2].text == "Test Bullet"  # Bullet added


def test_save(tmp_path, report):
    """Tests the save method."""
    report.save(tmp_path / "test.docx")
    assert os.path.exists(tmp_path / "test.docx")  # File saved
